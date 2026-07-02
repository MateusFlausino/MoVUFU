"""Gera o relatorio tecnico do MoV UFU baseado exclusivamente em OpenStreetMap."""

from __future__ import annotations

import json
import math
import re
from collections import defaultdict
from datetime import date
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATA_PATH = ROOT / "data" / "osm-santa-monica-accessibility.geojson"
OUTPUT_DIR = ROOT / "output"
ASSET_DIR = OUTPUT_DIR / "report-assets"
OUTPUT_PATH = OUTPUT_DIR / "Relatorio_MoV_UFU_OpenStreetMap.docx"
MAP_PATH = ASSET_DIR / "rede_osm_campus_santa_monica.png"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(32, 55, 72)
GRAY = RGBColor(90, 98, 106)
LIGHT_FILL = "F4F6F9"
TABLE_FILL = "E8EEF5"
WHITE = RGBColor(255, 255, 255)


def set_run_font(run, size=11, bold=None, italic=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{margin}"))
        if element is None:
            element = OxmlElement(f"w:{margin}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    table_width = table_pr.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_pr.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths_dxa)))
    table_width.set(qn("w:type"), "dxa")
    table_indent = table_pr.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_pr.append(table_indent)
    table_indent.set(qn("w:w"), "120")
    table_indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_hyperlink(paragraph, text, url):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.extend([properties, text_element])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def configure_page(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("MoV UFU | Relatorio tecnico OpenStreetMap")
    set_run_font(run, size=8.5, color=GRAY, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    run = footer.add_run("Pagina ")
    set_run_font(run, size=8.5, color=GRAY)
    add_page_field(footer)


def paragraph(document, text="", bold_lead=None, style=None, align=None):
    p = document.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        body = p.add_run(text[len(bold_lead):])
        set_run_font(body)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def create_list_numbering(document, kind):
    numbering = document.part.numbering_part.element
    abstract_ids = [int(item.get(qn("w:abstractNumId"))) for item in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))]
    attribute_name = f"_osm_{kind}_abstract_id"
    abstract_id = getattr(document, attribute_name, None)
    num_id = max(num_ids, default=0) + 1
    if abstract_id is None:
        abstract_id = max(abstract_ids, default=0) + 1
        setattr(document, attribute_name, abstract_id)
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        nsid = OxmlElement("w:nsid")
        nsid.set(qn("w:val"), f"A5{abstract_id:06X}"[-8:])
        multi_level = OxmlElement("w:multiLevelType")
        multi_level.set(qn("w:val"), "singleLevel")
        abstract.extend([nsid, multi_level])
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        number_format = OxmlElement("w:numFmt")
        number_format.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
        level_text = OxmlElement("w:lvlText")
        level_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
        justification = OxmlElement("w:lvlJc")
        justification.set(qn("w:val"), "left")
        paragraph_properties = OxmlElement("w:pPr")
        indentation = OxmlElement("w:ind")
        indentation.set(qn("w:left"), "540")
        indentation.set(qn("w:hanging"), "280")
        paragraph_properties.append(indentation)
        level.extend([start, number_format, level_text, justification, paragraph_properties])
        abstract.append(level)
        numbering.append(abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    number.append(abstract_reference)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    number.append(level_override)
    numbering.append(number)
    return num_id


def add_list_paragraph(document, num_id, text=""):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    if text:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_list_items(document, items, kind):
    num_id = create_list_numbering(document, kind)
    for item in items:
        p = add_list_paragraph(document, num_id, item)
        p.paragraph_format.space_after = Pt(4)


def add_numbered_items(document, items):
    add_list_items(document, items, "decimal")


def add_bulleted_items(document, items):
    add_list_items(document, items, "bullet")


def apply_portuguese_accents(document):
    replacements = {
        "Relatorio": "Relatório", "relatorio": "relatório", "tecnico": "técnico",
        "Tecnico": "Técnico", "acessivel": "acessível", "Acessivel": "Acessível",
        "Monica": "Mônica", "integracao": "integração", "reconhecimento": "reconhecimento",
        "Concepcao": "Concepção", "concepcao": "concepção", "aplicacao": "aplicação",
        "navegacao": "navegação", "geografica": "geográfica", "atualizavel": "atualizável",
        "correcoes": "correções", "operacao": "operação", "apresentacao": "apresentação",
        "travessia": "travessia", "sao": "são", "contem": "contém", "vertices": "vértices",
        "vertice": "vértice", "arestas": "arestas", "restricoes": "restrições",
        "publica": "pública", "portugues": "português", "brasileiro": "brasileiro",
        "primaria": "primária", "dados geograficos": "dados geográficos", "poligono": "polígono",
        "atributos": "atributos", "Caracteristicas": "Características", "fisicas": "físicas",
        "transicao": "transição", "Informacao": "Informação", "avaliacao": "avaliação",
        "sensorial": "sensorial", "aquisicao": "aquisição", "preparacao": "preparação",
        "geometrias": "geometrias", "vocabulario": "vocabulário", "fisico": "físico",
        "formula": "fórmula", "distancias": "distâncias", "superficie": "superfície",
        "proximo": "próximo", "profundidade": "profundidade", "oferecido": "oferecido",
        "Algoritmo": "Algoritmo", "mantem": "mantém", "seleciona": "seleciona",
        "ate": "até", "distancia": "distância", "usuario": "usuário", "metodologico": "metodológico",
        "parametros": "parâmetros", "usuarios": "usuários", "auditorias": "auditorias",
        "aplicacao": "aplicação", "cartografica": "cartográfica", "mosaicos": "mosaicos",
        "sobreposicao": "sobreposição", "interacao": "interação", "Distribuicao": "Distribuição",
        "numeros": "números", "edicao": "edição", "transcricao": "transcrição",
        "pontuacao": "pontuação", "contracoes": "contrações", "limitacoes": "limitações",
        "omissoes": "omissões", "desatualizados": "desatualizados", "predio": "prédio",
        "nao": "não", "certificacao": "certificação", "deficiencia": "deficiência",
        "atualizacao": "atualização", "registro fotografico": "registro fotográfico",
        "Conclusao": "Conclusão", "migracao": "migração", "consistente": "consistente",
        "atualizacoes": "atualizações", "comunitarias": "comunitárias", "representacao": "representação",
        "tecnicamente": "tecnicamente", "adequada": "adequada", "evolucao": "evolução",
        "participacao": "participação", "validacao": "validação", "Referencias": "Referências",
        "biblioteca": "biblioteca", "pagina": "página", "cartograficos": "cartográficos",
        "Caixeiro": "Caixeiro", "Chines": "Chinês", "exigencia": "exigência",
        "funcao": "função", "menor": "menor", "pontos": "pontos", "conectividade": "conectividade",
        "prototipo": "protótipo", "pesquisa": "pesquisa", "institucional": "institucional",
        "possui": "possui", "componente": "componente", "conectados": "conectados",
        "limite": "limite", "contidos": "contidos", "Conversao": "Conversão",
        "Uberlandia": "Uberlândia", "enfase": "ênfase", "mobilidade": "mobilidade",
        "dependencia": "dependência", "teoria": "teoria", "grafo": "grafo",
        "delimitacao": "delimitação", "selecao": "seleção", "reprodutivel": "reprodutível",
        "existencia": "existência", "associacao": "associação", "e colaborativo": "é colaborativo",
        "e calculado": "é calculado", "e convertido": "é convertido", "e executado": "é executado",
        "primaria e o": "primária é o",
    }
    patterns = sorted(replacements, key=len, reverse=True)

    def replace_text(text):
        for source in patterns:
            text = re.sub(rf"(?<![A-Za-z_]){re.escape(source)}(?![A-Za-z_])", replacements[source], text)
        return text

    containers = [document]
    for section in document.sections:
        containers.extend([section.header, section.footer])
    for container in containers:
        for p in container.paragraphs:
            for run in p.runs:
                for text_node in run._r.findall(".//" + qn("w:t")):
                    text_node.text = replace_text(text_node.text or "")
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            for text_node in run._r.findall(".//" + qn("w:t")):
                                text_node.text = replace_text(text_node.text or "")


def request_field_update(document):
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_callout(document, label, text):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p_pr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_FILL)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "6")
        border.set(qn("w:color"), "AAB4BE")
        borders.append(border)
    p_pr.append(borders)
    run = p.add_run(f"{label}: ")
    set_run_font(run, bold=True, color=DARK_BLUE)
    run = p.add_run(text)
    set_run_font(run)


def add_metrics_table(document, metrics):
    rows = [
        ("Caminhos de pedestres", str(metrics["paths"])),
        ("Travessias/pontos de acessibilidade", str(metrics["access_points"])),
        ("Destinos nomeados na fonte", str(metrics["destinations_total"])),
        ("Vertices do grafo", str(metrics["vertices"])),
        ("Arestas do grafo", str(metrics["edges"])),
        ("Componentes conectados", str(metrics["components"])),
        ("Vertices no maior componente", str(metrics["largest_component_vertices"])),
        ("Destinos conectados operacionais", str(metrics["connected_destinations"])),
    ]
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    headers = ("Indicador", "Resultado")
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade_cell(cell, TABLE_FILL)
        run = cell.paragraphs[0].add_run(value)
        set_run_font(run, bold=True, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in cells:
            for run in cell.paragraphs[0].runs:
                set_run_font(run)
    set_table_geometry(table, [6480, 2880])
    return table


def read_metrics(collection):
    def key(coordinate):
        return f"{coordinate[0]:.7f},{coordinate[1]:.7f}"

    graph = defaultdict(set)
    coordinates_by_key = {}
    edge_count = 0
    for feature in collection["features"]:
        if feature["properties"].get("feature_class") != "pedestrian_path":
            continue
        coordinates = feature["geometry"]["coordinates"]
        for left, right in zip(coordinates, coordinates[1:]):
            left_key, right_key = key(left), key(right)
            coordinates_by_key[left_key] = left
            coordinates_by_key[right_key] = right
            graph[left_key].add(right_key)
            graph[right_key].add(left_key)
            edge_count += 1

    components = []
    visited = set()
    for start in coordinates_by_key:
        if start in visited:
            continue
        component = {start}
        queue = [start]
        visited.add(start)
        while queue:
            current = queue.pop()
            for neighbor in graph[current]:
                if neighbor not in visited:
                    visited.add(neighbor)
                    component.add(neighbor)
                    queue.append(neighbor)
        components.append(component)
    largest = max(components, key=len)

    def haversine(left, right):
        radius = 6371000
        lat1, lat2 = math.radians(left[1]), math.radians(right[1])
        dlat = lat2 - lat1
        dlng = math.radians(right[0] - left[0])
        value = math.sin(dlat / 2) ** 2 + math.cos(lat1) * math.cos(lat2) * math.sin(dlng / 2) ** 2
        return 2 * radius * math.asin(math.sqrt(value))

    connected_destinations = 0
    for feature in collection["features"]:
        if feature["properties"].get("feature_class") != "destination":
            continue
        coordinate = feature["geometry"]["coordinates"]
        nearest = min(coordinates_by_key, key=lambda item: haversine(coordinate, coordinates_by_key[item]))
        connected_destinations += int(nearest in largest)

    metadata = collection["metadata"]
    return {
        "paths": metadata["pedestrian_path_count"],
        "access_points": metadata["accessibility_point_count"],
        "destinations_total": metadata["destination_count"],
        "vertices": len(coordinates_by_key),
        "edges": edge_count,
        "components": len(components),
        "largest_component_vertices": len(largest),
        "connected_destinations": connected_destinations,
    }


def create_map_figure(collection):
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    width, height = 1800, 1040
    margin = 70
    title_height = 80
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    all_coordinates = []
    for feature in collection["features"]:
        geometry = feature["geometry"]
        if geometry["type"] == "LineString":
            all_coordinates.extend(geometry["coordinates"])
        elif geometry["type"] == "Point":
            all_coordinates.append(geometry["coordinates"])
    min_lng = min(point[0] for point in all_coordinates)
    max_lng = max(point[0] for point in all_coordinates)
    min_lat = min(point[1] for point in all_coordinates)
    max_lat = max(point[1] for point in all_coordinates)
    map_width = width - 2 * margin
    map_height = height - title_height - 2 * margin

    def project(coordinate):
        x = margin + (coordinate[0] - min_lng) / (max_lng - min_lng) * map_width
        y = title_height + margin + (max_lat - coordinate[1]) / (max_lat - min_lat) * map_height
        return int(x), int(y)

    for feature in collection["features"]:
        properties = feature["properties"]
        geometry = feature["geometry"]
        if properties.get("feature_class") == "pedestrian_path":
            coordinates = geometry["coordinates"]
            color = "#dc2626" if properties.get("highway") == "steps" else "#0f766e"
            draw.line([project(point) for point in coordinates], fill=color, width=5, joint="curve")
        elif properties.get("feature_class") == "accessibility_point":
            x, y = project(geometry["coordinates"])
            draw.ellipse((x - 8, y - 8, x + 8, y + 8), fill="#f97316", outline="white", width=3)
        elif properties.get("feature_class") == "destination":
            x, y = project(geometry["coordinates"])
            draw.ellipse((x - 4, y - 4, x + 4, y + 4), fill="#2563eb")
    try:
        title_font = ImageFont.truetype("arial.ttf", 34)
    except OSError:
        title_font = ImageFont.load_default()
    title = "Rede pedestre OpenStreetMap - Campus Santa Mônica"
    title_box = draw.textbbox((0, 0), title, font=title_font)
    draw.text(((width - (title_box[2] - title_box[0])) / 2, 24), title, fill="#203748", font=title_font)
    image.save(MAP_PATH, optimize=True)


def build_document(collection, metrics):
    document = Document()
    configure_styles(document)
    configure_page(document)

    for _ in range(5):
        document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("RELATORIO TECNICO")
    set_run_font(run, size=11, bold=True, color=BLUE)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("MoV UFU")
    set_run_font(run, size=30, bold=True, color=NAVY)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("Roteamento acessivel no Campus Santa Monica com OpenStreetMap")
    set_run_font(run, size=15, color=DARK_BLUE)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Modelagem do grafo, algoritmo de menor caminho e integracao com reconhecimento de voz")
    set_run_font(run, size=10.5, italic=True, color=GRAY)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(90)
    run = p.add_run(f"Uberlandia - MG | {date.today().strftime('%d/%m/%Y')}")
    set_run_font(run, size=11, bold=True, color=NAVY)
    document.add_page_break()

    document.add_heading("Resumo executivo", level=1)
    paragraph(document, "Este relatorio apresenta a reformulacao do MoV UFU como uma aplicacao de navegacao pedestre baseada exclusivamente em dados do OpenStreetMap (OSM). A planta DWG/PDF deixa de compor a fonte operacional do sistema. Os caminhos, travessias e destinos sao obtidos da base colaborativa OSM, convertidos para GeoJSON e transformados em um grafo ponderado sobre o qual e executado o algoritmo de Dijkstra.")
    paragraph(document, "A versao analisada contem 117 caminhos de pedestres, 32 pontos de travessia ou acessibilidade e 98 destinos nomeados. A rede possui 382 vertices, 398 arestas e seis componentes conectados. Para impedir a apresentacao de rotas inexistentes, a aplicacao opera sobre o maior componente, com 356 vertices e 81 destinos conectados.")
    add_callout(document, "Resultado principal", "A aplicacao passou a calcular rotas diretamente sobre a geometria OpenStreetMap. Escadas e trechos explicitamente marcados como wheelchair=no recebem penalidade elevada, enquanto trechos limitados ou sem pavimentacao recebem custo adicional.")

    document.add_heading("1. Concepcao do projeto", level=1)
    paragraph(document, "O MoV UFU foi concebido para apoiar o deslocamento de pedestres no Campus Santa Monica da Universidade Federal de Uberlandia, com enfase em pessoas com mobilidade reduzida. A proposta combina cartografia colaborativa, teoria dos grafos, processamento geoespacial e uma interface web responsiva.")
    paragraph(document, "A reformulacao substitui a dependencia do desenho tecnico por uma fonte geografica aberta e continuamente atualizavel. Essa decisao permite que correcoes feitas no OpenStreetMap sejam incorporadas ao sistema sem reeditar ou reexportar um arquivo CAD.")
    add_callout(document, "Nota de nomenclatura", "Embora o diretorio do projeto mencione o Problema do Carteiro Viajante, a funcao implementada e o menor caminho entre uma origem e um destino. Nao se trata do Problema do Caixeiro Viajante nem do Carteiro Chines, pois nao ha exigencia de visitar todos os pontos ou percorrer todas as arestas.")

    document.add_heading("2. Objetivos", level=1)
    objective_items = (
        "Representar os caminhos de pedestres do campus como um grafo georreferenciado.",
        "Calcular o menor percurso entre blocos e equipamentos nomeados no OSM.",
        "Evitar ou desestimular escadas e trechos com restricoes conhecidas de acessibilidade.",
        "Permitir atualizacao reprodutivel a partir da base publica do OpenStreetMap.",
        "Disponibilizar selecao por listas, texto e comandos de voz em portugues brasileiro.",
    )
    add_bulleted_items(document, objective_items)

    document.add_heading("3. Fonte de dados e delimitacao", level=1)
    paragraph(document, "A fonte primaria e o OpenStreetMap, projeto colaborativo de dados geograficos abertos. A coleta abrange o limite cadastrado da Universidade Federal de Uberlandia no Campus Santa Monica. O importador consulta a API OSM, identifica o poligono universitario e conserva apenas elementos localizados dentro desse limite.")
    paragraph(document, "Os dados sao armazenados em GeoJSON com tres classes: pedestrian_path para caminhos lineares, accessibility_point para travessias e elementos de acessibilidade, e destination para blocos, bibliotecas, equipamentos e outros locais nomeados.")

    document.add_heading("4. Convencoes de mapeamento no OpenStreetMap", level=1)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for index, value in enumerate(("Elemento", "Tag principal", "Uso no sistema")):
        cell = table.rows[0].cells[index]
        shade_cell(cell, TABLE_FILL)
        run = cell.paragraphs[0].add_run(value)
        set_run_font(run, bold=True, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    rows = [
        ("Caminho pedestre", "highway=footway/path/pedestrian", "Forma as arestas navegaveis."),
        ("Travessia", "highway=crossing", "Ponto de verificacao de rampa e seguranca."),
        ("Meio-fio rebaixado", "kerb=lowered ou kerb=flush", "Evidencia de transicao potencialmente acessivel."),
        ("Escada", "highway=steps", "Recebe penalidade elevada."),
        ("Acesso por cadeira de rodas", "wheelchair=yes/limited/no", "Modifica o custo do trecho."),
        ("Piso tatil", "tactile_paving=yes/no", "Informacao para avaliacao de acessibilidade sensorial."),
        ("Caracteristicas fisicas", "surface, width, incline, smoothness", "Base para evoluir o modelo de custo."),
    ]
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value
            for run in cell.paragraphs[0].runs:
                set_run_font(run, size=9.5)
    set_table_geometry(table, [2300, 3000, 4060])

    document.add_heading("5. Processo de aquisicao e preparacao", level=1)
    steps = [
        "Consultar a API publica do OpenStreetMap na caixa geografica do campus.",
        "Localizar o poligono com amenity=university e nome Universidade Federal de Uberlandia.",
        "Filtrar caminhos, travessias, atributos de acessibilidade e destinos contidos no campus.",
        "Converter geometrias e atributos para uma FeatureCollection GeoJSON.",
        "Construir o grafo, identificar componentes conectados e selecionar a maior rede operacional.",
        "Atualizar a interface, o vocabulário de voz e o cache do aplicativo.",
    ]
    add_numbered_items(document, steps)

    create_map_figure(collection)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    inline_shape = p.add_run().add_picture(str(MAP_PATH), width=Inches(6.35))
    inline_shape._inline.docPr.set("descr", "Mapa esquemático da rede pedestre do Campus Santa Mônica, com caminhos em verde, travessias em laranja e destinos em azul.")
    inline_shape._inline.docPr.set("title", "Rede pedestre OpenStreetMap do Campus Santa Mônica")
    caption = document.add_paragraph("Figura 1 - Rede pedestre extraida do OpenStreetMap. Linhas verdes representam caminhos, laranja representa travessias e azul representa destinos nomeados.")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    for run in caption.runs:
        set_run_font(run, size=8.5, italic=True, color=GRAY)

    document.add_heading("6. Modelagem do grafo", level=1)
    paragraph(document, "Cada coordenada consecutiva de uma LineString e convertida em vertice. Coordenadas iguais, arredondadas a sete casas decimais, compartilham a mesma chave e estabelecem a conectividade entre vias. Cada par consecutivo forma uma aresta bidirecional.")
    paragraph(document, "O comprimento fisico da aresta e calculado pela formula de Haversine, apropriada para distancias sobre a superficie terrestre. Os destinos nomeados sao associados ao vertice mais proximo. Antes do roteamento, uma busca em profundidade identifica os componentes da rede; apenas o maior componente e oferecido como rede operacional.")

    document.add_heading("7. Algoritmo de roteamento", level=1)
    paragraph(document, "O menor caminho e calculado pelo algoritmo de Dijkstra. Para cada vertice, o algoritmo mantem o menor custo conhecido desde a origem, seleciona o vertice nao visitado de menor custo e relaxa as arestas adjacentes ate atingir o destino ou esgotar a componente.")
    paragraph(document, "O custo de uma aresta corresponde ao seu comprimento multiplicado por um fator de acessibilidade. Trechos comuns usam fator 1. Trechos com wheelchair=limited ou surface=unpaved usam fator 4. Escadas e trechos com wheelchair=no usam fator 1000. A distancia exibida ao usuario permanece a distancia fisica; a penalidade serve apenas para escolher a rota.")
    add_callout(document, "Cuidado metodologico", "Os fatores 4 e 1000 sao heuristicas de projeto, nao parametros normativos. Eles devem ser calibrados com usuarios, auditorias de campo e requisitos de acessibilidade antes de uso institucional.")

    document.add_heading("8. Arquitetura da aplicacao", level=1)
    for label, text in (
        ("Camada de dados", "GeoJSON gerado pelo importador Python a partir do OpenStreetMap."),
        ("Camada cartografica", "Leaflet com mosaicos do OpenStreetMap e sobreposicao vetorial dos caminhos."),
        ("Camada de grafo", "JavaScript no navegador para vertices, arestas, componentes e Dijkstra."),
        ("Camada de interacao", "Seletores de origem/destino, formulario textual e Web Speech API em pt-BR."),
        ("Distribuicao", "Aplicacao web progressiva com manifesto e service worker para recursos locais."),
    ):
        paragraph(document, f"{label}: {text}", bold_lead=f"{label}:")

    document.add_heading("9. Resultados atuais", level=1)
    paragraph(document, "A tabela resume o estado da extracao realizada em 1 de julho de 2026. Os numeros podem variar a cada nova edicao da comunidade OSM.")
    add_metrics_table(document, metrics)
    results_note = paragraph(document, "A existencia de seis componentes demonstra que alguns caminhos ainda nao compartilham vertices, mesmo quando parecem proximos visualmente. O bloco 1E, por exemplo, foi identificado em uma componente pequena na extracao analisada; por isso nao deve ser conectado artificialmente ao bloco 3E. A correcao adequada e completar a geometria no OpenStreetMap apos verificacao local.")
    results_note.paragraph_format.space_before = Pt(8)

    document.add_heading("10. Reconhecimento de voz", level=1)
    paragraph(document, "O reconhecimento usa a Web Speech API configurada para portugues brasileiro. O sistema considera ate cinco alternativas de transcricao, resultados intermediarios e um limite de 15 segundos. Nomes e referencias dos destinos OSM compoem o vocabulario dinamico.")
    paragraph(document, "O analisador normaliza acentos, pontuacao, numeros falados e contracoes como de, do e da. Frases como 'do bloco 1A para o bloco 3E' sao decompostas em origem e destino, desde que ambos estejam conectados a rede operacional.")

    document.add_heading("11. Validacao, riscos e limitacoes", level=1)
    limitation_items = (
        "O OpenStreetMap e colaborativo e pode conter omissoes, geometrias desconectadas ou atributos desatualizados.",
        "Uma travessia mapeada nao comprova a existencia de rampa, meio-fio rebaixado ou piso tatil.",
        "A associacao de um predio ao caminho mais proximo nao substitui o mapeamento da entrada correta.",
        "Trechos sem tags de acessibilidade sao tratados como neutros, e nao como comprovadamente acessiveis.",
        "O sistema deve ser validado em campo e com pessoas com deficiencia antes de orientar deslocamentos reais.",
    )
    add_bulleted_items(document, limitation_items)

    document.add_heading("12. Protocolo recomendado de atualizacao", level=1)
    protocol_steps = (
        "Realizar levantamento em campo com GPS e registro fotografico autorizado.",
        "Editar caminhos e atributos no editor iD, StreetComplete ou JOSM.",
        "Garantir que caminhos que se cruzam compartilhem um vertice quando houver passagem fisica.",
        "Registrar kerb, wheelchair, tactile_paving, incline, surface e width somente quando confirmados.",
        "Executar tools/import_osm_santa_monica.py para atualizar o GeoJSON.",
        "Executar os testes e revisar visualmente as rotas antes da publicacao.",
    )
    add_numbered_items(document, protocol_steps)

    document.add_heading("13. Conclusao", level=1)
    paragraph(document, "A migracao para OpenStreetMap torna o MoV UFU geograficamente consistente, aberto a atualizacoes comunitarias e independente de reexportacoes CAD. A representacao em grafo e o algoritmo de Dijkstra fornecem uma base tecnicamente adequada para rotas ponto a ponto. O principal trabalho futuro e melhorar a completude dos atributos de acessibilidade e a conectividade da rede por meio de levantamento de campo.")
    paragraph(document, "A versao atual deve ser entendida como prototipo de pesquisa e apoio ao mapeamento, nao como certificacao de rota acessivel. Sua evolucao depende da qualidade dos dados, da participacao dos usuarios e da validacao institucional.")

    document.add_heading("Referencias", level=1)
    references = [
        ("OpenStreetMap - pagina principal e dados cartograficos", "https://www.openstreetmap.org/"),
        ("OpenStreetMap Wiki - Key:wheelchair", "https://wiki.openstreetmap.org/wiki/Key:wheelchair"),
        ("OpenStreetMap Wiki - Key:kerb", "https://wiki.openstreetmap.org/wiki/Key:kerb"),
        ("OpenStreetMap Wiki - Key:ramp", "https://wiki.openstreetmap.org/wiki/Key:ramp"),
        ("OpenStreetMap Wiki - Key:tactile_paving", "https://wiki.openstreetmap.org/wiki/Key:tactile_paving"),
        ("OpenStreetMap Wiki - Overpass Turbo", "https://wiki.openstreetmap.org/wiki/Overpass_turbo"),
        ("Leaflet - biblioteca de mapas web", "https://leafletjs.com/"),
    ]
    reference_num_id = create_list_numbering(document, "bullet")
    for title, url in references:
        p = add_list_paragraph(document, reference_num_id)
        add_hyperlink(p, title, url)

    apply_portuguese_accents(document)
    request_field_update(document)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)


def main():
    collection = json.loads(DATA_PATH.read_text(encoding="utf-8"))
    metrics = read_metrics(collection)
    build_document(collection, metrics)
    print(json.dumps({"output": str(OUTPUT_PATH), "metrics": metrics}, ensure_ascii=False))


if __name__ == "__main__":
    main()
